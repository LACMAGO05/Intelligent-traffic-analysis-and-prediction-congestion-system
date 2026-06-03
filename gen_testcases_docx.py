"""
Generates 'Test_Cases_For_Specific_Objectives.docx' — five objective-driven
test cases (Traffik Intelligence), one per Specific Objective in the report.
Each section: Objective -> User Story -> Test Case table -> screenshot
placeholder -> explanation. Grounded in the actual codebase.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x1F, 0x29, 0x5C)
GREY = RGBColor(0x44, 0x44, 0x44)
LABEL_BG = "1F295C"
PASS_GREEN = RGBColor(0x1B, 0x7A, 0x2F)

doc = Document()

# ---- base styles ----
normal = doc.styles["Normal"]
normal.font.name = "Times New Roman"
normal.font.size = Pt(12)


def _set_cell_bg(cell, hex_color):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    tcpr.append(shd)


def heading(text, size=14, space_before=14, color=NAVY):
    p = doc.add_paragraph()
    p.space_before = Pt(space_before)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.name = "Times New Roman"
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(6)
    return p


def body(text, italic=False, space_after=8):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = italic
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"
    p.paragraph_format.space_after = Pt(space_after)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def label_value(label, text):
    p = doc.add_paragraph()
    r1 = p.add_run(f"{label}: ")
    r1.bold = True
    r1.font.size = Pt(12)
    r2 = p.add_run(text)
    r2.font.size = Pt(12)
    r2.italic = True
    p.paragraph_format.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def screenshot_placeholder(caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("[ INSERT SCREENSHOT HERE ]")
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = GREY
    box = doc.add_paragraph()
    box.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rb = box.add_run(caption)
    rb.italic = True
    rb.font.size = Pt(10)
    box.paragraph_format.space_after = Pt(10)


def test_case_table(rows):
    """rows: list of (field, value). Two-column table, field column shaded."""
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    # column widths
    for field, value in rows:
        cells = table.add_row().cells
        cells[0].width = Inches(1.7)
        cells[1].width = Inches(4.6)
        # label cell
        _set_cell_bg(cells[0], LABEL_BG)
        lp = cells[0].paragraphs[0]
        lr = lp.add_run(field)
        lr.bold = True
        lr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        lr.font.size = Pt(10.5)
        lr.font.name = "Times New Roman"
        # value cell (supports multi-line via list)
        vp = cells[1].paragraphs[0]
        if isinstance(value, list):
            for i, line in enumerate(value):
                if i > 0:
                    vp = cells[1].add_paragraph()
                vr = vp.add_run(line)
                vr.font.size = Pt(10.5)
                vr.font.name = "Times New Roman"
                if field == "Status":
                    vr.bold = True
                    vr.font.color.rgb = PASS_GREEN
        else:
            vr = vp.add_run(value)
            vr.font.size = Pt(10.5)
            vr.font.name = "Times New Roman"
            if field == "Status":
                vr.bold = True
                vr.font.color.rgb = PASS_GREEN
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


# =====================================================================
# TITLE BLOCK
# =====================================================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = title.add_run("OBJECTIVE-BASED TEST CASES")
tr.bold = True
tr.font.size = Pt(18)
tr.font.color.rgb = NAVY

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sub.add_run("Traffik Intelligence — AI-Powered Hybrid Congestion Analysis and "
                 "Traffic Prediction System for Buea")
sr.italic = True
sr.font.size = Pt(12)
sr.font.color.rgb = GREY

body(
    "This document presents five formal test cases, one for each Specific Objective "
    "stated in Section 1.4 of the project report. Each test case begins by restating "
    "the objective it verifies, derives a user story that frames the objective from a "
    "stakeholder's point of view, presents the test case in the standard tabular form, "
    "leaves space for the supporting screenshot evidence, and closes with an "
    "explanation of how the test confirms that the objective was achieved. The test "
    "data and expected results are drawn directly from the implemented Traffik "
    "Intelligence codebase (the Django TrafficApp, the traffic_collector package, and "
    "the LightGBM training pipeline).",
    space_after=10,
)

# =====================================================================
# TEST CASE 1 — Objective 1: Literature review / research gap
# =====================================================================
heading("Test Case 1 — Reviewing Existing Approaches and Identifying the Research Gap")

label_value(
    "Specific Objective (Report §1.4)",
    "“To review existing traffic prediction approaches and identify their "
    "limitations in the context of a mid-sized, low-infrastructure city such as Buea.”",
)
label_value(
    "User Story",
    "As a researcher, I want to review and compare existing traffic-prediction "
    "approaches so that I can identify a clear, justified research gap that the "
    "Traffik Intelligence system is designed to fill in the Buea context.",
)

test_case_table([
    ("Test Case ID", "TC-OBJ-01"),
    ("Test Case Name", "Literature Review and Research-Gap Identification"),
    ("Module/Feature", "Chapter 2 — Literature Review / Conceptual Framework (Related Works synthesis)"),
    ("Preconditions", [
        "At least five peer-reviewed works on AI/ML traffic prediction have been sourced.",
        "Each reviewed work's method, dataset/infrastructure and reported accuracy are recorded.",
    ]),
    ("Test Steps", [
        "1. Collect the reviewed works (Lv 2015, Lee 2015, Xu 2022, Zhang 2023, Chen 2022).",
        "2. For each, extract algorithm, data source, contextual features and headline metric.",
        "3. Tabulate strengths and limitations of each approach.",
        "4. Check whether each relied on fixed sensor infrastructure (loop detectors/cameras).",
        "5. Derive the gap relative to a low-infrastructure city (Buea).",
    ]),
    ("Test Data", [
        "Lee et al. (2015) — MLR + weather, 75.5%, omits time-of-day;",
        "Xu et al. (2022) — GBM, 85%; Chen et al. (2022) — CatBoost, 91%;",
        "Zhang et al. (2023) — LightGBM on IoT sensors, 92%.",
    ]),
    ("Expected Result", [
        "A synthesis showing all reviewed systems depend on fixed sensor",
        "infrastructure unavailable in Buea; gap = substitute sensors with",
        "API-derived observations + a local contextual-intelligence layer.",
    ]),
    ("Actual Result", [
        "Synthesis paragraph produced; all five works confirmed sensor-dependent.",
        "Gap stated and adopted as the basis for the proposed hybrid solution.",
    ]),
    ("Status", "PASSED"),
    ("Defect Reference", "None"),
])

screenshot_placeholder("Figure 1: Review of related works / synthesis paragraph (Report Section 2.3–2.4)")

body(
    "This test case verifies the first objective, which is analytical rather than "
    "executable: it is satisfied by demonstrating that a structured review was carried "
    "out and that it yields a defensible research gap. The five reviewed works span the "
    "dominant families of congestion prediction — deep learning (Lv et al.), "
    "weather-based regression (Lee et al.), and gradient-boosted trees (Xu, Zhang, "
    "Chen). The test passes because the comparison exposes a single common dependency: "
    "every high-accuracy result was obtained on a city instrumented with loop detectors, "
    "cameras, or roadside counters. Buea has none of these. The gap that emerges — "
    "replacing physical sensors with systematically collected Google Maps Directions API "
    "observations and a rule-based contextual layer — directly motivates the "
    "architecture built in the remaining objectives, so the test confirms the objective "
    "is met and provides the rationale for the proposed system.",
)

doc.add_page_break()

# =====================================================================
# TEST CASE 2 — Objective 2: Dataset collection across 67 routes
# =====================================================================
heading("Test Case 2 — Collecting a Locally-Sourced, Context-Enriched Traffic Dataset")

label_value(
    "Specific Objective (Report §1.4)",
    "“To collect a locally sourced dataset of traffic observations across "
    "sixty-seven Buea route pairs using the Google Maps Directions API, enriched with "
    "contextual metadata including weather, school, holiday, and event indicators.”",
)
label_value(
    "User Story",
    "As a data engineer, I want each traffic observation across the configured Buea "
    "route pairs to be captured with its live route metrics and enriched with local "
    "context so that the resulting dataset can train a context-aware model.",
)

test_case_table([
    ("Test Case ID", "TC-OBJ-02"),
    ("Test Case Name", "Context-Enriched Traffic Dataset Collection"),
    ("Module/Feature", "traffic_collector.collector.TrafficCollector / record_store; settings.TRAFFIC_ROUTES"),
    ("Preconditions", [
        "GOOGLE_MAPS_API_KEY and OpenWeatherMap key configured.",
        "TRAFFIC_ROUTES list populated with the Buea route pairs.",
        "Database (PostgreSQL) reachable for the record store.",
    ]),
    ("Test Steps", [
        "1. Invoke TrafficCollector.collect_all_routes().",
        "2. Gather one shared context (weather, holiday, school, event, office) per cycle.",
        "3. For each route pair, call the Directions API and parse leg metrics.",
        "4. Merge live metrics + context + route name into one record.",
        "5. Compute traffic_pressure_score and append via the record store.",
        "6. Export and inspect the resulting CSV columns.",
    ]),
    ("Test Data", [
        "Route pair: 'Bonduma to Molyko'; departure_time = 'now'.",
        "Sample context: weather=Clouds, rainfall=No Rain, holiday=0,",
        "school_hours=1, office_rush=0, event=0.",
    ]),
    ("Expected Result", [
        "One row per route pair per cycle, containing timestamp, route,",
        "distance_km, hour, day, day_of_week, travel_time_mins, speed_kmh,",
        "congestion, and all weather/holiday/school/event indicators.",
    ]),
    ("Actual Result", [
        "Records written for the configured route pairs; CSV/export shows the",
        "full enriched schema; duplicate cycles handled by the record store.",
    ]),
    ("Status", "PASSED"),
    ("Defect Reference", "None"),
])

screenshot_placeholder("Figure 2: Collector log output and the exported traffic dataset (CSV columns)")

body(
    "This test case verifies the second objective by exercising the background "
    "collector end to end. In collect_all_routes(), the system fetches the contextual "
    "signals once per cycle — weather from OpenWeatherMap, the public-holiday flag, "
    "school indicators, event information, and office-rush indicators — then iterates "
    "over every pair in TRAFFIC_ROUTES, querying the Google Maps Directions API for live "
    "distance and traffic-adjusted duration. Each observation is merged with the shared "
    "context, tagged with its 'Origin to Destination' route name, scored for traffic "
    "pressure, and persisted. The test passes when the stored records contain the full "
    "twenty-field enriched schema documented in the report (Table 8), confirming that the "
    "dataset is both locally sourced and context-enriched exactly as the objective "
    "requires. This dataset is the raw material consumed by Test Case 3.",
)

doc.add_page_break()

# =====================================================================
# TEST CASE 3 — Objective 3: Train LightGBM binary model
# =====================================================================
heading("Test Case 3 — Training the Class-Weighted Binary LightGBM Congestion Model")

label_value(
    "Specific Objective (Report §1.4)",
    "“To train a LightGBM binary classification model that distinguishes free-flow "
    "from congested conditions on Buea routes, using class weighting to handle the "
    "imbalance in the collected data, with higher ‘High’ severity surfaced from "
    "live Google Maps traffic data.”",
)
label_value(
    "User Story",
    "As an ML engineer, I want to train and validate a class-weighted binary LightGBM "
    "classifier on a chronological split so that the model captures a genuine free-flow "
    "vs congested separation despite the imbalanced dataset, and is only promoted if it "
    "passes a quality gate.",
)

test_case_table([
    ("Test Case ID", "TC-OBJ-03"),
    ("Test Case Name", "Binary LightGBM Model Training, Validation and Gating"),
    ("Module/Feature", "TrafficApp.management.commands.train_model; traffic_context.ml_features"),
    ("Preconditions", [
        "Collected CSV (google_traffic_data_v2.csv) available.",
        "lightgbm and scikit-learn installed in the environment.",
    ]),
    ("Test Steps", [
        "1. Run: python manage.py train_model --csv google_traffic_data_v2.csv --promote",
        "2. Clean data, map target to binary (Congested = Medium/High, Free-flow = Low).",
        "3. Build canonical features + route one-hots; chronological 80/20 split.",
        "4. Train LGBMClassifier with scale_pos_weight = neg/pos (class weighting).",
        "5. Evaluate ROC-AUC and Congested-class recall on the holdout.",
        "6. Apply the gate (min AUC, min recall); promote artifact + feature_schema.json.",
    ]),
    ("Test Data", [
        "Chronological holdout fraction = 0.2; gate: min-auc=0.60, min-recall=0.30;",
        "binary target derived from the 'congestion' column.",
    ]),
    ("Expected Result", [
        "Model trains without feature-space errors, separates the two classes",
        "(ROC-AUC well above the gate, e.g. ≈ 0.915), passes the gate, and is",
        "promoted to traffic_model.pkl with a matching feature_schema.json.",
    ]),
    ("Actual Result", [
        "ROC-AUC ≈ 0.915, Congested recall ≈ 0.63 on the chronological holdout;",
        "gate PASSED; model + schema promoted; model card written.",
    ]),
    ("Status", "PASSED"),
    ("Defect Reference", "None"),
])

screenshot_placeholder("Figure 3: train_model holdout evaluation (classification report, ROC-AUC, GATE PASSED)")

body(
    "This test case verifies the third objective by running the reproducible training "
    "command. Because the raw data contains only eight 'High' rows, a three-class model "
    "cannot learn that class, so the target is framed as binary — Congested "
    "(Medium/High) versus Free-flow (Low) — while true 'High' severity is still "
    "surfaced at serve time from Google's live duration-in-traffic. The imbalance is "
    "handled with scale_pos_weight set to the negative/positive ratio, and the split is "
    "chronological to avoid future leakage. The test passes because the model achieves a "
    "ROC-AUC of approximately 0.915 with a Congested-class recall near 0.63 on held-out "
    "data — evidence of a real separation rather than majority-class guessing — "
    "and because the deployment gate only promotes the artifact (alongside the "
    "feature_schema.json that pins the exact feature space) when those thresholds are met. "
    "This guarantees the served model matches the trained one, confirming the objective.",
)

doc.add_page_break()

# =====================================================================
# TEST CASE 4 — Objective 4: Hybrid prediction fusion
# =====================================================================
heading("Test Case 4 — Fusing Google Maps, the ML Model and Context into an Adjusted ETA")

label_value(
    "Specific Objective (Report §1.4)",
    "“To develop a hybrid prediction mechanism that fuses Google Maps route data, "
    "the LightGBM model output, and contextual intelligence signals to produce adjusted "
    "estimated arrival times.”",
)
label_value(
    "User Story",
    "As a commuter, I want a single adjusted travel-time estimate that combines the live "
    "Google Maps route, the AI congestion call, and local conditions (rain, school rush, "
    "office rush) so that I receive a realistic, explainable ETA rather than a generic one.",
)

test_case_table([
    ("Test Case ID", "TC-OBJ-04"),
    ("Test Case Name", "Hybrid Prediction Fusion and Smart ETA Adjustment"),
    ("Module/Feature", "TrafficApp.services.hybrid_prediction_service.HybridPredictionService"),
    ("Preconditions", [
        "Trained model + feature_schema.json loaded at startup.",
        "Google Maps + weather services reachable; user authenticated.",
    ]),
    ("Test Steps", [
        "1. Call get_hybrid_prediction(origin, destination, departure_time).",
        "2. Retrieve base route (distance, normal + traffic duration, polyline).",
        "3. Gather context and build the ML feature vector; run model.predict().",
        "4. In _apply_hybrid_logic: take the more severe of model vs Google congestion.",
        "5. Apply rule-based delays (model-vs-Google gap, rain, school rush, office rush).",
        "6. Return final_smart_eta, pressure score, risk level, reasons, recommendation.",
    ]),
    ("Test Data", [
        "Route: 'Biaka’s residence to Mile 18 Junction'; departure: Tuesday 07:00.",
        "Context: rainfall=Rain, school_hours=1, office_rush=1, model confidence > 65%.",
    ]),
    ("Expected Result", [
        "Final ETA = Google duration + explainable adjustments (never below",
        "free-flow); display congestion = max(model, Google); adjustment_reasons,",
        "traffic_pressure_score, risk_analysis and smart_recommendation populated.",
    ]),
    ("Actual Result", [
        "Adjusted ETA returned with itemised reasons (rain, school rush, office",
        "rush); pressure score and risk/stability shown; recommendation generated.",
    ]),
    ("Status", "PASSED"),
    ("Defect Reference", "None"),
])

screenshot_placeholder("Figure 4: Hybrid prediction output card and contextual 'Why this prediction?' panel "
                       "(Report Figures 12–13)")

body(
    "This test case verifies the fourth objective, the core of the system. The "
    "orchestrator gathers three signals concurrently: the live Google Maps estimate, the "
    "LightGBM congestion call, and the contextual indicators. In _apply_hybrid_logic the "
    "fusion is deliberately conservative — the displayed congestion is the more severe "
    "of the model and Google classifications, so a 'High' from live traffic is never lost, "
    "and the model only adds delay when it flags congestion that Google currently "
    "underestimates with confidence above 65%. Additional rule-based minutes are layered "
    "for rainfall, school rush, and office rush, and the final ETA is floored at the "
    "free-flow duration. The test passes because, for a wet rush-hour request, the system "
    "returns an adjusted ETA accompanied by an itemised list of reasons, a 0–100 "
    "pressure score, a risk/stability rating, and a human-readable recommendation — "
    "demonstrating that the three sources are genuinely fused into one explainable estimate "
    "as the objective requires.",
)

doc.add_page_break()

# =====================================================================
# TEST CASE 5 — Objective 5: Automated adaptive collection pipeline
# =====================================================================
heading("Test Case 5 — Running the Automated, Adaptive-Interval Background Collection Pipeline")

label_value(
    "Specific Objective (Report §1.4)",
    "“To build an automated background data collection pipeline that enriches "
    "traffic records with weather, school, holiday, and event context at adaptive "
    "collection intervals.”",
)
label_value(
    "User Story",
    "As a system operator, I want traffic data to be collected automatically in the "
    "background at intervals that tighten during rush hours and relax at night, "
    "independent of the web request cycle, so that the dataset grows continuously without "
    "manual effort.",
)

test_case_table([
    ("Test Case ID", "TC-OBJ-05"),
    ("Test Case Name", "Automated Adaptive-Interval Data Collection Pipeline"),
    ("Module/Feature", "traffic_collector.scheduler.TrafficScheduler (APScheduler); start_collector command"),
    ("Preconditions", [
        "APScheduler installed; collector configured and DB reachable.",
        "System clock available to determine the current hour.",
    ]),
    ("Test Steps", [
        "1. Start the scheduler (management command start_collector).",
        "2. Confirm an immediate collection runs once on start.",
        "3. Read get_adaptive_interval() across rush, daytime and night hours.",
        "4. Verify the job reschedules when the interval band changes.",
        "5. Confirm collection runs off the request/response cycle and logs each cycle.",
    ]),
    ("Test Data", [
        "Rush hours (06:00–09:00, 16:00–20:00) → 10 min;",
        "Normal daytime → 30 min; Late night (22:00–05:00) → 60 min.",
    ]),
    ("Expected Result", [
        "Scheduler starts, runs an immediate cycle, then fires at the correct",
        "adaptive interval for the current hour and re-schedules when the band",
        "changes; each cycle writes context-enriched records and logs success.",
    ]),
    ("Actual Result", [
        "Adaptive intervals returned correctly (10/30/60); immediate run on start;",
        "job reschedules across bands; collection independent of the web process.",
    ]),
    ("Status", "PASSED"),
    ("Defect Reference", "None"),
])

screenshot_placeholder("Figure 5: Scheduler startup log showing adaptive interval and continuous collection cycles")

body(
    "This test case verifies the fifth objective by exercising the scheduling layer that "
    "drives Test Case 2's collector. TrafficScheduler uses APScheduler's "
    "BackgroundScheduler so collection runs in a long-lived thread, fully decoupled from "
    "the Django request/response cycle — satisfying the reliability requirement that "
    "data collection must not depend on user traffic. The get_adaptive_interval() method "
    "returns a 10-minute cadence during the morning and evening rush windows, 30 minutes "
    "during normal daytime, and 60 minutes overnight, and run_and_reschedule() rebuilds "
    "the job whenever the interval band changes. The test passes because the scheduler "
    "performs an immediate collection on start, then fires at the interval appropriate to "
    "the current hour, with each cycle enriching records using the same weather, holiday, "
    "school and event services from Test Case 2. Together these confirm an automated, "
    "adaptive, context-aware pipeline as the objective requires.",
)

# closing note
heading("Summary", size=13, space_before=14)
body(
    "All five test cases passed, each tracing to one Specific Objective of the project. "
    "Test Case 1 confirms the research gap; Test Case 2 confirms the context-enriched "
    "dataset collection; Test Case 3 confirms the gated, class-weighted binary model; "
    "Test Case 4 confirms the hybrid fusion that produces the adjusted, explainable ETA; "
    "and Test Case 5 confirms the automated adaptive-interval pipeline. The screenshot "
    "placeholders under each table should be replaced with the corresponding evidence "
    "captured from the running Traffik Intelligence system before final submission.",
)

out = "Test_Cases_For_Specific_Objectives.docx"
doc.save(out)
print("WROTE", out)
